import aiohttp
from typing import List, Tuple
from docx import Document
import io

async def download_docx_from_url(url: str) -> bytes:
    """
    Download a DOCX file from the given URL into memory.
    Raises on non-200 HTTP.
    """
    async with aiohttp.ClientSession() as session:
        async with session.get(url) as resp:
            if resp.status != 200:
                raise RuntimeError(f"Failed to fetch DOCX ({resp.status})")
            return await resp.read()

def extract_text_from_docx(docx_bytes: bytes) -> List[Tuple[int, str]]:
    """
    Extracts text from each paragraph of the DOCX.
    Returns list of (paragraph_number, text).
    """
    paragraphs: List[Tuple[int, str]] = []
    
    # Load document from bytes
    doc = Document(io.BytesIO(docx_bytes))
    
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text:  # Only add non-empty paragraphs
            paragraphs.append((i + 1, text))
    
    return paragraphs

def extract_tables_from_docx(docx_bytes: bytes) -> List[dict]:
    """
    Extracts tables from DOCX with their structure.
    Returns list of table dictionaries.
    """
    tables = []
    doc = Document(io.BytesIO(docx_bytes))
    
    for table_idx, table in enumerate(doc.tables):
        table_data = {
            "table_index": table_idx,
            "rows": []
        }
        
        for row_idx, row in enumerate(table.rows):
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_data["rows"].append(row_data)
        
        tables.append(table_data)
    
    return tables
